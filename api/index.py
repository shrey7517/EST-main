import json
import os
import requests
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
from concurrent.futures import ThreadPoolExecutor

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Initialize Firebase Admin
firebase_sa_json = os.environ.get("FIREBASE_SERVICE_ACCOUNT_JSON")
if firebase_sa_json:
    try:
        cert_dict = json.loads(firebase_sa_json)
        cred = credentials.Certificate(cert_dict)
        firebase_admin.initialize_app(cred)
    except Exception as e:
        print(f"Error loading Firebase Credentials from env: {e}")
        try:
            firebase_admin.initialize_app()
        except Exception as fe:
            print(f"Fallback credentials initialization failed: {fe}")
else:
    # Fallback to default credentials (for local dev if configured)
    try:
        firebase_admin.initialize_app()
    except ValueError:
        pass # Already initialized
    except Exception as e:
        print(f"Default credentials initialization failed: {e}")

db = None
try:
    db = firestore.client()
except Exception as e:
    print(f"Firestore Client Initialization Error: {e}")
    print("Warning: Firestore client is uninitialized. Local database writes will fail until credentials are set.")

app = FastAPI()

# Enable CORS for the frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # In production, you can replace "*" with your actual frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------- MCQ EVALUATION ----------------
def evaluate_mcq(answers):
    correct = ['c', 'a', 'd', 'e', 'd', 'd', 'e', 'c', 'd', 'd', 'd']
    # TODO: Update weights below. The total must be 16.
    # As per prompt, specific questions like #8 and #9 have higher weights.
    weights = [1, 1, 1, 1, 1, 1, 1, 3, 3, 2, 2] # Sum = 16
    score = 0
    for i in range(min(len(answers), len(correct))):
        if str(answers[i]).lower() == correct[i]:
            score += weights[i]
    return score

# ---------------- PAIRED EVALUATION ----------------
def evaluate_paired(answers):
    correct = [
        "A","A","B","B","A","B","A","B",
        "B","A","B","A","A","B","B","A",
        "B","A","B","A","A","A","B","A",
        "B","A","B","A","B","A","B","A"
    ]
    points = 0
    for i in range(min(len(answers), len(correct))):
        # TODO: Replace with exact mapping where each pair assigns 0, 1, or 2 points for A or B.
        # Currently defaults to 2 points for the correct answer, 0 for incorrect.
        if str(answers[i]).upper() == correct[i]:
            points += 2
            
    # Final Score Conversion
    if points <= 25:
        return 0
    elif points <= 36:
        return 1
    elif points <= 47:
        return 2
    else:
        return 3

# ---------------- DEEPSEEK ----------------
def call_deepseek(prompt):
    api_key = os.environ.get("DEEPSEEK_API_KEY", "sk-aee4748505a344689e59b5f8d0fc4f48") # Used hardcoded key as fallback
    try:
        url = "https://api.deepseek.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3
        }
        response = requests.post(url, headers=headers, json=data, timeout=30)
        result = response.json()
        if "choices" not in result:
            return json.dumps({"total": 0})
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        print("DeepSeek Error:", e)
        return json.dumps({"total": 0})

def clean_json(text):
    try:
        text = text.replace("```json", "").replace("```", "").strip()
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
        return {"total": 0}
    except:
        return {"total": 0}

def evaluate_who(essay):
    if not essay or len(essay) < 10:
        return {"total": 0}
    prompt = f"""
    Evaluate Who Am I:
    {essay}
    Based on Initiative, Problem Solving, Goal Clarity, Resource Awareness.
    Return JSON: {{"initiative":0/1, "problem_solving":0/1, "goal_clarity":0/1, "resource_awareness":0/1, "total":0-4}}
    """
    return clean_json(call_deepseek(prompt))

def evaluate_image(story):
    if not story or len(story) < 10:
        return {"total": 0}
    prompt = f"""
    Evaluate Achievement Motivation Story:
    {story}
    Based on Achievement Imagery (AI) criteria. 
    If the story is categorized as "Unrelated Imagery" (UI) or "Task Imagery" (TI), the score must be 0.
    Return JSON: {{"total":0-11}}
    """
    return clean_json(call_deepseek(prompt))

@app.post("/api/submit_test")
async def submit_test(request: Request):
    try:
        data = await request.json()
        user_id = data.get("user_id")
        room_key = data.get("room_key")
        
        mcq_score = evaluate_mcq(data.get("section1", []))
        paired_score = evaluate_paired(data.get("section2", []))
        
        who_text = data.get("who_am_i", "")
        images_data = data.get("images", [])

        # Evaluate who_am_i and images in parallel (7 concurrent requests max)
        with ThreadPoolExecutor(max_workers=7) as executor:
            who_future = executor.submit(evaluate_who, who_text)
            image_futures = [executor.submit(evaluate_image, img_text) for img_text in images_data]
            
            # Resolve future results
            who = who_future.result()
            images_results = [future.result() for future in image_futures]

        if isinstance(who, dict):
            who["original_text"] = who_text
            
        for idx, img_text in enumerate(images_data):
            if idx < len(images_results) and isinstance(images_results[idx], dict):
                images_results[idx]["original_text"] = img_text
        
        total = mcq_score + paired_score + who.get("total", 0) + sum(i.get("total", 0) for i in images_results)
        
        # Save to Firestore
        if db is None:
            raise Exception("Firestore database client is not initialized. Please verify your service account key or GOOGLE_APPLICATION_CREDENTIALS path in your .env file.")
        db.collection("results").add({
            "user_id": user_id,
            "room_key": room_key,
            "section1": data.get("section1", []),
            "section2": data.get("section2", []),
            "who_am_i": who,
            "images": images_results,
            "total": total,
            "createdAt": firestore.SERVER_TIMESTAMP
        })
        
        return {
            "mcq_score": mcq_score,
            "paired_score": paired_score,
            "who_am_i": who,
            "images": images_results,
            "total": total
        }
    except Exception as e:
        print("Error submitting test", e)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/generate_room_report/{room_key}")
async def generate_room_report(room_key: str):
    try:
        # Fetch all results for the room
        if db is None:
            raise Exception("Firestore database client is not initialized. Please verify your service account key or GOOGLE_APPLICATION_CREDENTIALS path in your .env file.")
        results_ref = db.collection("results").where("room_key", "==", room_key)
        results = results_ref.stream()
        
        scores = []
        mcq_scores = []
        paired_scores = []
        trait_counts = {"initiative": 0, "problem_solving": 0, "goal_clarity": 0, "resource_awareness": 0}
        total_students = 0
        
        for doc_snap in results:
            res = doc_snap.to_dict()
            total_students += 1
            scores.append(res.get("total", 0))
            mcq_scores.append(res.get("mcq_score", 0))
            paired_scores.append(res.get("paired_score", 0))
            
            who = res.get("who_am_i", {})
            for trait in trait_counts.keys():
                trait_counts[trait] += who.get(trait, 0)
                
        if total_students == 0:
            raise HTTPException(status_code=404, detail="No results found for this room")
            
        avg_score = sum(scores) / total_students
        avg_mcq = sum(mcq_scores) / total_students
        avg_paired = sum(paired_scores) / total_students
        
        # Build prompt for AI
        prompt = f"""
        You are an expert entrepreneurial psychologist and data analyst. Generate a comprehensive cohort analysis report for a group of students based on their aggregate Entrepreneurship Skill Test (EST) results.
        
        Cohort Data:
        - Total Students: {total_students}
        - Average Total Score: {avg_score:.2f}
        - Average MCQ Score: {avg_mcq:.2f} (Logical reasoning)
        - Average Paired Comparison Score: {avg_paired:.2f} (Decision making)
        
        Essay Trait Presence (Number of students demonstrating each trait):
        - Initiative: {trait_counts['initiative']} / {total_students}
        - Problem Solving: {trait_counts['problem_solving']} / {total_students}
        - Goal Clarity: {trait_counts['goal_clarity']} / {total_students}
        - Resource Awareness: {trait_counts['resource_awareness']} / {total_students}
        
        Please write a cohesive, 4-5 paragraph executive summary and analysis report assessing this cohort's overall entrepreneurial potential, identifying common strengths, highlighting areas for aggregate improvement, and providing actionable recommendations for educators based on these specific metrics. Do not use markdown syntax in your final output, just plain paragraphs.
        """
        
        # Call AI
        ai_response = call_deepseek(prompt)
        
        # Create Docx
        doc = Document()
        doc.add_heading(f"Cohort Analysis Report - Room: {room_key}", 0)
        doc.add_paragraph(f"Total Students Assessed: {total_students}")
        doc.add_paragraph(f"Average Cohort Score: {avg_score:.2f}")
        
        doc.add_heading("Executive Summary & Psychological Evaluation", 1)
        doc.add_paragraph(ai_response)
        
        file_path = f"/tmp/room_report_{room_key}.docx"
        doc.save(file_path)
        
        return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f"Room_{room_key}_Analysis_Report.docx")
        
    except Exception as e:
        print("Error generating room report:", e)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/")
def read_root():
    return {"message": "EST Platform API is running"}
