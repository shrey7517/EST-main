from flask import Flask, request, jsonify
from flask_cors import CORS
import sqlite3, json, os, requests
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
from docx import Document

# ---------------- APP ----------------
app = Flask(__name__, static_folder='static')
CORS(app)

# ---------------- DATABASE ----------------
def init_db():
    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    c.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT UNIQUE,
        password TEXT,
        role TEXT
    )
    """)

    c.execute("""
    CREATE TABLE IF NOT EXISTS results (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        section1 TEXT,
        section2 TEXT,
        who_am_i TEXT,
        images TEXT,
        total INTEGER
    )
    """)

    conn.commit()
    conn.close()

# ---------------- MCQ EVALUATION ----------------
def evaluate_mcq(answers):
    # 🔥 Replace with actual answer key
    correct = ['c', 'a', 'd', 'e', 'd', 'd', 'e', 'c', 'd', 'd', 'd']

    score = 0
    for i in range(min(len(answers), len(correct))):
        if answers[i] == correct[i]:
            score += 1

    return score

# ---------------- PAIRED EVALUATION ----------------
def evaluate_paired(answers):
    # 🔥 Replace with official answer key if available
    correct = [
        "A","A","B","B","A","B","A","B",
        "B","A","B","A","A","B","B","A",
        "B","A","B","A","A","A","B","A",
        "B","A","B","A","B","A","B","A"
    ]

    score = 0
    for i in range(min(len(answers), len(correct))):
        if answers[i] == correct[i]:
            score += 1

    return score

# ---------------- DEEPSEEK ----------------
def call_deepseek(prompt):
    try:
        url = "https://api.deepseek.com/v1/chat/completions"

        headers = {
            "Authorization": f"Bearer {os.environ.get('DEEPSEEK_API_KEY')}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3
        }

        response = requests.post(url, headers=headers, json=data)
        result = response.json()

        if "choices" not in result:
            return json.dumps({"total": 0})

        return result["choices"][0]["message"]["content"]

    except Exception as e:
        print("DeepSeek Error:", e)
        return json.dumps({"total": 0})

def clean_json(text):
    try:
        text = text.replace("```json", "").replace("```", "")
        return json.loads(text)
    except:
        return {"total": 0}

# ---------------- EVALUATION ----------------
def evaluate_who(essay):
    prompt = f"""
    Evaluate Who Am I:

    {essay}

    Based on:
    - Initiative
    - Problem Solving
    - Goal Clarity
    - Resource Awareness

    Return JSON:
    {{
        "initiative":0/1,
        "problem_solving":0/1,
        "goal_clarity":0/1,
        "resource_awareness":0/1,
        "total":0-4
    }}
    """
    return clean_json(call_deepseek(prompt))

def evaluate_image(story):
    prompt = f"""
    Evaluate Achievement Motivation Story:

    {story}

    Return JSON:
    {{
        "total":0-11
    }}
    """
    return clean_json(call_deepseek(prompt))

# ---------------- REPORT ----------------
def generate_report(user_id, mcq_score, paired_score, who, images, total):
    doc = Document()

    doc.add_heading("EST Report", 0)
    doc.add_paragraph(f"User ID: {user_id}")
    doc.add_paragraph(f"MCQ Score: {mcq_score}")
    doc.add_paragraph(f"Paired Score: {paired_score}")
    doc.add_paragraph(f"Total Score: {total}")

    doc.add_heading("Who Am I", 1)
    doc.add_paragraph(json.dumps(who, indent=2))

    doc.add_heading("Image Evaluation", 1)
    for i, img in enumerate(images):
        doc.add_paragraph(f"Image {i+1}: {img}")

    filename = f"report_{user_id}.docx"
    doc.save(filename)

    return filename

# ---------------- AUTH ----------------
@app.route("/register", methods=["POST"])
def register():
    data = request.json

    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    try:
        c.execute("""
        INSERT INTO users (name,email,password,role)
        VALUES (?,?,?,?)
        """, (
            data["name"],
            data["email"],
            generate_password_hash(data["password"]),
            "student"
        ))
        conn.commit()
    except:
        return jsonify({"error": "User already exists"})

    return jsonify({"message": "Registered successfully"})

@app.route("/login", methods=["POST"])
def login():
    data = request.json

    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    user = c.execute("SELECT * FROM users WHERE email=?", (data["email"],)).fetchone()

    if user and check_password_hash(user[3], data["password"]):
        return jsonify({
            "id": user[0],
            "name": user[1],
            "role": user[4]
        })

    return jsonify({"error": "Invalid credentials"})

# ---------------- SUBMIT TEST ----------------
@app.route("/submit", methods=["POST"])
def submit():
    data = request.json

    mcq_score = evaluate_mcq(data.get("section1", []))
    paired_score = evaluate_paired(data.get("section2", []))
    who_text = data.get("who_am_i", "")
    who = evaluate_who(who_text)
    if isinstance(who, dict):
        who["original_text"] = who_text
        
    images = []
    for i in data.get("images", []):
        res = evaluate_image(i)
        if isinstance(res, dict):
            res["original_text"] = i
        images.append(res)

    total = (
        mcq_score +
        paired_score +
        who.get("total", 0) +
        sum(i.get("total", 0) for i in images)
    )

    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    c.execute("""
    INSERT INTO results (user_id,section1,section2,who_am_i,images,total)
    VALUES (?,?,?,?,?,?)
    """, (
        data["user_id"],
        json.dumps(data.get("section1", [])),
        json.dumps(data.get("section2", [])),
        json.dumps(who),
        json.dumps(images),
        total
    ))

    conn.commit()
    conn.close()

    report_file = generate_report(
        data["user_id"],
        mcq_score,
        paired_score,
        who,
        images,
        total
    )

    return jsonify({
        "mcq_score": mcq_score,
        "paired_score": paired_score,
        "who_am_i": who,
        "images": images,
        "total": total,
        "report": report_file
    })

# ---------------- ADMIN ----------------
@app.route("/admin/results", methods=["GET"])
def admin_results():
    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    data = c.execute("SELECT * FROM results").fetchall()
    conn.close()

    return jsonify(data)

@app.route("/admin/edit", methods=["POST"])
def edit_score():
    data = request.json

    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    c.execute("UPDATE results SET total = ? WHERE id = ?", (data["score"], data["id"]))

    conn.commit()
    conn.close()

    return jsonify({"status": "updated"})

# ---------------- OWNER ----------------
@app.route("/owner/create-admin", methods=["POST"])
def create_admin():
    data = request.json

    conn = sqlite3.connect("est.db")
    c = conn.cursor()

    c.execute("""
    INSERT INTO users (name,email,password,role)
    VALUES (?,?,?,?)
    """, (
        data["name"],
        data["email"],
        generate_password_hash(data["password"]),
        "admin"
    ))

    conn.commit()
    conn.close()

    return jsonify({"message": "Admin created"})

# ---------------- EXPORT EXCEL ----------------
@app.route("/export-excel", methods=["GET"])
def export_excel():
    conn = sqlite3.connect("est.db")
    df = pd.read_sql_query("SELECT * FROM results", conn)

    file_path = "results.xlsx"
    df.to_excel(file_path, index=False)

    conn.close()

    return jsonify({"file": file_path})


# ---------------- RUN ----------------
if __name__ == "__main__":
    init_db()

    # SET YOUR KEY HERE
    os.environ["DEEPSEEK_API_KEY"] = "sk-aee4748505a344689e59b5f8d0fc4f48"

   port = int(os.environ.get("PORT", 5000))
app.run(host="0.0.0.0", port=port)
